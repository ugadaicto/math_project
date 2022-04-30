import docx
import random
from flask import Flask, url_for, render_template, redirect
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, BooleanField, SubmitField
from wtforms.validators import DataRequired
from forms.login_form import LoginForm
from forms.register_form import RegisterForm
from forms.quiz_form import QuizForm
import sqlite3


app = Flask(__name__)
app.config['SECRET_KEY'] = 'rabotai_pozhaluysta'
USER_LOGGED_IN = False
CURRENT_QUESTION = 0
ANSWERED_CORRECTLY = 0
LOGIN = None


def get_questions_from_database():
    con = sqlite3.connect('quiz.db')
    cur = con.cursor()
    questions = [question[0] for question in cur.execute(f"""SELECT question FROM questions""").fetchall()]
    return questions


QUESTIONS = random.choices(get_questions_from_database(), k=10)


def check_answer(question, answer):
    con = sqlite3.connect('quiz.db')
    cur = con.cursor()
    correct_answer = cur.execute(f"""SELECT correct_answer FROM questions WHERE question='{question}'""").fetchall()[0][
        0]
    return answer == correct_answer, correct_answer


def password_is_correct(user_name, password):
    con = sqlite3.connect('profiles.db')
    cur = con.cursor()
    result = cur.execute(f"""SELECT password FROM profiles WHERE user_name = '{user_name}'""")
    result = result.fetchall()
    if result:
        return result[0][0] == password
    return False


def make_new_profile(user_name, password):
    con = sqlite3.connect('profiles.db')
    cur = con.cursor()
    cur.execute(
        f"""INSERT INTO profiles(user_name, password, last_result, best_result)
         VALUES('{user_name}', '{password}', 0, 0)""")
    con.commit()


def get_best_result_from_database(user_name):
    con = sqlite3.connect('profiles.db')
    cur = con.cursor()
    best_result = cur.execute(
        f"""SELECT best_result FROM profiles WHERE user_name = '{user_name}'""").fetchall()
    print(f'best res: {best_result[0][0]}')
    if best_result[0][0]:
        return best_result[0][0]
    return 0


def get_last_result_from_database(user_name):
    con = sqlite3.connect('profiles.db')
    cur = con.cursor()
    last_result = cur.execute(
        f"""SELECT last_result FROM profiles WHERE user_name = '{user_name}'""").fetchall()
    if last_result[0][0]:
        return last_result[0][0]
    return 0


def insert_results_into_database(new_best, result, user_name):
    con = sqlite3.connect('profiles.db')
    cur = con.cursor()
    if new_best:
        cur.execute(f"""UPDATE profiles SET best_result = {result} WHERE user_name = '{user_name}'""")
    cur.execute(f"""UPDATE profiles SET last_result = {result} WHERE user_name = '{user_name}'""")
    con.commit()


def lessons(file_name):
    c = docx.Document(file_name)
    v = ''
    for i in c.paragraphs[2:]:
        v += i.text + '\n'
    return [c.paragraphs[0].text, v]


@app.route('/login', methods=['GET', 'POST'])
def login():
    global USER_LOGGED_IN
    global LOGIN
    form = LoginForm()
    if form.validate_on_submit():
        login, password = form.username._value(), form.password._value()

        if password_is_correct(login, password):
            USER_LOGGED_IN = True
            LOGIN = login
            return render_template('menu.html', user_logged_in=USER_LOGGED_IN)

        else:
            return render_template('log.html', title='Авторизация', form=form, custom_error="Неверный логин или пароль")

    return render_template('log.html', title='Авторизация', form=form)


@app.route('/register', methods=['GET', 'POST'])
def register():
    global USER_LOGGED_IN
    global LOGIN
    form = RegisterForm()
    if form.validate_on_submit():
        user_name, password, password_check = form.username._value(), form.password._value(), form.password_check._value()

        if password != password_check:
            return render_template('reg.html', title='Регистрация', form=form,
                                   custom_error="Введнные пароли не сходятся")

        make_new_profile(user_name, password)
        USER_LOGGED_IN = True
        LOGIN = user_name
        return render_template('menu.html', user_logged_in=USER_LOGGED_IN)

    return render_template('reg.html', title='Регистрация', form=form)


@app.route("/quiz", methods=['GET', 'POST'])
def quiz():
    global LOGIN
    global CURRENT_QUESTION
    global ANSWERED_CORRECTLY
    global USER_LOGGED_IN
    global QUESTIONS
    if USER_LOGGED_IN:
        form = QuizForm()
        questions = QUESTIONS
        print(questions)
        question = questions[CURRENT_QUESTION]
        if form.validate_on_submit():
            chosen_answer = None
            if form.data["move_on"]:
                CURRENT_QUESTION += 1

                if CURRENT_QUESTION == len(questions):
                    best_result = get_best_result_from_database(LOGIN)

                    insert_results_into_database(best_result < ANSWERED_CORRECTLY * 10, ANSWERED_CORRECTLY, LOGIN)

                    print(USER_LOGGED_IN)
                    print(ANSWERED_CORRECTLY, ANSWERED_CORRECTLY / len(questions) * 100)
                    return render_template('success.html', result=int(ANSWERED_CORRECTLY / len(questions) * 100))

                question = questions[CURRENT_QUESTION]
                return render_template('quiz.html', question=question, form=form)

            if form.data["answer1"]:
                print(CURRENT_QUESTION)
                chosen_answer = 1
                response, correct_answer = check_answer(question, 1)
            if form.data["answer2"]:
                chosen_answer = 2
                response, correct_answer = check_answer(question, 2)
            if form.data["answer3"]:
                chosen_answer = 3
                response, correct_answer = check_answer(question, 3)
            if form.data["answer4"]:
                chosen_answer = 4
                response, correct_answer = check_answer(question, 4)

            if response:
                ANSWERED_CORRECTLY += 1
                print(ANSWERED_CORRECTLY)

            question = questions[CURRENT_QUESTION]
            print(question)

            return render_template('quiz.html', question=question, form=form,
                                   correct_answer=correct_answer, chosen_answer=chosen_answer, answered=True)
    else:
        return render_template('menu.html', user_logged_in=USER_LOGGED_IN)

    question = questions[CURRENT_QUESTION]

    return render_template('quiz.html', question=question, form=form)


@app.route('/')
def menu():
    global USER_LOGGED_IN
    return render_template('menu.html', user_logged_in=USER_LOGGED_IN)


@app.route("/lesson_<int:number>/")
def lesson(number):
    c = lessons(f'conspects/lesson_{number}.docx')
    return render_template("lesson.html", title=c[0], lesson_text=c[1])


if __name__ == '__main__':
    app.run(port=8080, host='127.0.0.1')
